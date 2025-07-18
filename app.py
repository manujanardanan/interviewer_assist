import streamlit as st
from openai import OpenAI
import json
import traceback
import unicodedata
from st_audiorec import st_audiorec
from fpdf import FPDF
import io
from docx import Document

# --- Page Configuration ---
st.set_page_config(page_title="GenAI Interview Agent", page_icon="🤖", layout="wide")
st.title("🤖 GenAI Interview Agent")

# --- OpenAI Client Initialization ---
try:
    client = OpenAI(api_key=st.secrets["openai"]["api_key"])
except Exception:
    st.error("OpenAI API key not found. Please add it to your Streamlit secrets.", icon="🚨")
    st.stop()

# --- PDF Generation Function ---
def create_pdf(details, report_data):
    def sanitize_text(text):
        # This function robustly handles special characters for PDF generation.
        return unicodedata.normalize('NFKD', str(text)).encode('ascii', 'ignore').decode('ascii')

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, f"Interview Report for: {sanitize_text(details['name'])}", 0, 1, 'C')
    pdf.ln(5)
    pdf.set_font("Arial", '', 12)
    pdf.cell(0, 10, f"Role Level: {details['role_level']} | Salary Expectation: {details['lpa']} LPA", 0, 1, 'C')
    pdf.ln(10)

    for i, item in enumerate(report_data):
        pdf.set_font("Arial", 'B', 12)
        pdf.multi_cell(0, 5, f"Question {i+1}: {sanitize_text(item['question'])}")
        pdf.ln(2)
        
        pdf.set_font("Arial", 'B', 11)
        pdf.multi_cell(0, 5, "Candidate's Answer:")
        pdf.set_font("Arial", 'I', 11)
        pdf.multi_cell(0, 5, f"{sanitize_text(item.get('answer', 'N/A'))}")
        pdf.ln(2)

        pdf.set_font("Arial", 'B', 11)
        pdf.multi_cell(0, 5, "Evaluation:")
        pdf.set_font("Arial", '', 11)
        
        eval_data = item.get('evaluation', {}).get('evaluation', {})
        if eval_data:
            clarity = eval_data.get('clarity', {})
            correctness = eval_data.get('correctness', {})
            depth = eval_data.get('depth', {})
            pdf.multi_cell(0, 5, f"  Clarity: {clarity.get('score', 0)}/10 - {sanitize_text(clarity.get('justification', ''))}")
            pdf.multi_cell(0, 5, f"  Correctness: {correctness.get('score', 0)}/10 - {sanitize_text(correctness.get('justification', ''))}")
            pdf.multi_cell(0, 5, f"  Depth: {depth.get('score', 0)}/10 - {sanitize_text(depth.get('justification', ''))}")
        
        summary = item.get('evaluation', {}).get('overall_summary', 'N/A')
        pdf.multi_cell(0, 5, f"Summary: {sanitize_text(summary)}")
        pdf.ln(8)
    
    return pdf.output(dest='S').encode('latin1')

# --- Session State Initialization ---
if 'status' not in st.session_state:
    st.session_state.status = 'setup'
    st.session_state.candidate_details = {}
    st.session_state.questions_to_ask = []
    st.session_state.notes = ""
    st.session_state.labeled_transcript = ""
    st.session_state.detailed_report = []
    st.session_state.audio_bytes = None
    st.session_state.question_number = 0

# --- Helper Functions ---
def start_new_interview():
    # Resets the entire session to start over
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    st.rerun()

def get_ai_response(prompt_text, model="gpt-4-turbo", as_json=False):
    """Generic function to call OpenAI API."""
    try:
        messages = [{"role": "system", "content": "You are a helpful assistant designed to output JSON if requested."}, {"role": "user", "content": prompt_text}]
        if as_json:
            response = client.chat.completions.create(model=model, response_format={"type": "json_object"}, messages=messages)
            return json.loads(response.choices[0].message.content)
        else:
            response = client.chat.completions.create(model=model, messages=messages, temperature=0.7, max_tokens=1000)
            return response.choices[0].message.content
    except Exception as e:
        st.error(f"AI Error: {e}")
        st.error(traceback.format_exc())
        return None

def generate_question(role_level, question_number):
    """Generates a question based on the interview stage."""
    prompt = ""
    if question_number == 1:
        prompt = f"Generate ONE open-ended, situational GenAI interview question for a '{role_level}' level candidate. The question should explore their experience with a challenging project. Return ONLY the question text."
    elif question_number == 2:
        prompt = f"Generate a problem-situation for a '{role_level}' level GenAI professional. Describe a scenario where a GenAI project is failing (e.g., high latency, poor accuracy, unexpected costs). Phrase the question to ask the candidate to identify potential causes. Example: 'A deployed model is suddenly hallucinating... what are the first things you would investigate?' Return ONLY the question text."
    elif question_number == 3:
        prompt = f"Generate a problem for a '{role_level}' level GenAI professional where the problem and diagnosis are given. The candidate must explain how to solve it. Example: 'Our RAG system is slow because of inefficient embedding lookups. How would you architect a solution to fix this?' Return ONLY the question text."
    elif question_number == 4:
        prompt = f"Generate a moderately tough diagnostic GenAI interview question for a '{role_level}' level candidate. The question should describe a specific, non-obvious technical issue in a deployed GenAI system and ask for a debugging process. Example: 'A RAG pipeline returns less relevant results for certain query types, but not all. The vector embeddings haven't changed. What specific components would you investigate first, and in what order?' Return ONLY the question text."
    
    if prompt:
        return get_ai_response(prompt)
    return "All questions have been asked."
    
def create_word_document(details, report_data):
    """Generates a .docx file from the report data."""
    document = Document()
    document.add_heading(f"Interview Report for: {details['name']}", level=1)
    document.add_paragraph(f"Role Level: {details['role_level']} | Salary Expectation: {details['lpa']} LPA")
    document.add_paragraph() # Spacer

    for i, item in enumerate(report_data):
        document.add_heading(f"Question {i+1}: {item['question']}", level=2)
        
        p_answer = document.add_paragraph()
        p_answer.add_run("Candidate's Answer: ").bold = True
        p_answer.add_run(item.get('answer', 'N/A')).italic = True
        
        eval_data = item.get('evaluation', {})
        summary = eval_data.get('overall_summary', 'No summary available.')
        scores = eval_data.get('evaluation', {})
        
        p_eval = document.add_paragraph()
        p_eval.add_run("Assessment/Evaluation: ").bold = True
        p_eval.add_run(summary)
        
        clarity = scores.get('clarity', {}).get('score', 0)
        correctness = scores.get('correctness', {}).get('score', 0)
        depth = scores.get('depth', {}).get('score', 0)
        avg_score = round((clarity + correctness + depth) / 3)

        p_score = document.add_paragraph()
        p_score.add_run("Score: ").bold = True
        p_score.add_run(f"{avg_score}/10")
        document.add_paragraph() # Spacer

    # Save the document to an in-memory stream
    doc_stream = io.BytesIO()
    document.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

# --- STAGE 1: SETUP ---
if st.session_state.status == 'setup':
    st.header("Stage 1: Candidate Details")
    with st.form("setup_form"):
        name = st.text_input("Candidate Name")
        lpa = st.number_input("Salary Expectation (LPA)", min_value=10, value=30)
        submitted = st.form_submit_button("Proceed to Question Prep")
        if submitted and name:
            st.session_state.candidate_details = {"name": name, "lpa": lpa, "role_level": "Senior" if lpa > 35 else "Mid"}
            st.session_state.status = 'question_prep'
            st.rerun()

# --- STAGE 2: QUESTION PREPARATION ---
elif st.session_state.status == 'question_prep':
    st.header("Stage 2: Prepare Interview Questions")
    st.info("Generate and refine your 4 questions before starting the recording.")

    if st.session_state.get('rephrase_triggered', False):
        st.session_state.rephrase_triggered = False 
        if st.session_state.questions_to_ask:
            last_question = st.session_state.questions_to_ask[-1]
            prompt = f"Rephrase the following interview question to be clearer or provide a different angle: '{last_question}'"
            with st.spinner("Rephrasing..."):
                rephrased_q = get_ai_response(prompt)
                if rephrased_q:
                    st.session_state.questions_to_ask[-1] = rephrased_q
    
    if 'question_number' in st.session_state and st.session_state.question_number > len(st.session_state.questions_to_ask):
        with st.spinner("Generating..."):
            new_question = generate_question(st.session_state.candidate_details['role_level'], st.session_state.question_number)
            if new_question and new_question != "All questions have been asked.":
                st.session_state.questions_to_ask.append(new_question)

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Question Controls")
        next_q_num = st.session_state.question_number + 1
        
        if next_q_num <= 4:
            st.button(f"Suggest Question {next_q_num}/4", on_click=lambda: st.session_state.update(question_number=st.session_state.question_number + 1))
        
        if st.session_state.questions_to_ask:
            st.button("Rephrase Last Question", on_click=lambda: st.session_state.update(rephrase_triggered=True))
    
    with col2:
        st.subheader("Prepared Questions")
        if not st.session_state.questions_to_ask:
            st.write("Click 'Suggest Question 1/4' to begin.")
        else:
            for i, q in enumerate(st.session_state.questions_to_ask):
                st.markdown(f"**{i+1}.** {q}")
    
    st.markdown("---")
    if st.button("Proceed to Live Recording", type="primary"):
        st.session_state.status = 'recording'
        st.rerun()

# --- STAGE 3: LIVE RECORDING ---
elif st.session_state.status == 'recording':
    st.header("Stage 3: Live Recording")
    st.success("RECORDING IN PROGRESS...")
    st.info("Ask the prepared questions and record the candidate's responses in one continuous session. Click 'Stop' when done.")

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Questions to Ask")
        if not st.session_state.questions_to_ask:
            st.warning("No questions were prepared.")
        else:
            for i, q in enumerate(st.session_state.questions_to_ask):
                st.markdown(f"**{i+1}.** {q}")
    
    with col2:
        st.subheader("Audio Recorder")
        audio_bytes = st_audiorec()
        
        # This now only saves the audio to the session state when the recording is stopped.
        # It does not automatically navigate away.
        if audio_bytes and len(audio_bytes) > 1000:
            st.session_state.audio_bytes = audio_bytes

    st.subheader("Interviewer's Notes")
    st.session_state.notes = st.text_area("Take live notes here:", height=200, value=st.session_state.notes)
    
    st.markdown("---")

    # --- NEW: Confirmation and Navigation Logic ---
    # This section now appears only AFTER audio has been recorded.
    if st.session_state.audio_bytes:
        # 1. Provide clear feedback that the recording is saved.
        st.success("✅ Recording complete and saved.")
        
        # 2. Provide an explicit button to move to the next step.
        if st.button("Proceed to Evaluation", type="primary"):
            st.session_state.status = 'processing'
            st.rerun()
            
# --- STAGE 4: PROCESSING & CONFIRMATION ---
elif st.session_state.status in ['processing', 'transcript_confirmation']:
    if st.session_state.status == 'processing':
        with st.spinner("Step 1/2: Transcribing audio... This may take a few minutes."):
            raw_transcript = ""
            try:
                if st.session_state.audio_bytes:
                    transcript_response = client.audio.transcriptions.create(model="whisper-1", file=("interview.wav", st.session_state.audio_bytes))
                    raw_transcript = transcript_response.text
                else:
                    st.error("No audio data found to process.")
                    st.session_state.status = 'recording'
                    st.stop()
            except Exception as e:
                st.error(f"Transcription Failed: {e}")
                st.session_state.status = 'recording'
                st.stop()
        
        with st.spinner("Step 2/2: AI is labeling speakers in the transcript..."):
            labeling_prompt = f"""You are an assistant that processes interview transcripts. Reformat the transcript below by adding 'Interviewer:' and 'Candidate:' labels.
            **CRITICAL RULES:**
            1. The candidate's speech directly follows each question.
            2. If the speech after a question is not a real answer (e.g., they just repeat the question or say "I don't know"), label it as 'Candidate:' but keep the content as is.
            3. **Do NOT invent or create any text for the candidate's answer.** If you cannot find a distinct response for a question, label the answer as 'Candidate: [No clear answer was recorded for this question]'.
            QUESTIONS ASKED: {st.session_state.questions_to_ask}
            FULL TRANSCRIPT: {raw_transcript}"""
            labeled_transcript = get_ai_response(labeling_prompt)
            if labeled_transcript:
                st.session_state.labeled_transcript = labeled_transcript
                st.session_state.status = 'transcript_confirmation'
                st.rerun()
            else:
                st.error("AI Speaker Labeling Failed.")
                st.session_state.status = 'recording'
    
    if st.session_state.status == 'transcript_confirmation':
        st.header("Stage 4: Confirm Speaker Labels")
        st.info("✅ **CRITICAL STEP:** Review the AI-labeled transcript below. **You can and should edit the text in this box** to correct any errors before running the final evaluation.")
        
        st.session_state.labeled_transcript = st.text_area("**Editable Labeled Transcript:**", value=st.session_state.labeled_transcript, height=400)
        
        if st.button("Confirm Transcript & Run Final Evaluation", type="primary"):
            st.session_state.status = 'evaluating'
            st.rerun()

# --- STAGE 5: FINAL REPORT ---
elif st.session_state.status in ['evaluating', 'report']:
    st.header(f"Stage 5: Evaluation & Final Report")
    
    # This block runs the evaluation pipeline once.
    if st.session_state.status == 'evaluating':
        report_data = []
        with st.spinner("Running detailed per-question evaluation..."):
            for i, question in enumerate(st.session_state.questions_to_ask):
                st.write(f"Evaluating answer for question {i+1}...")
                
                extract_prompt = f"""From the labeled transcript below, extract ONLY the 'Candidate:' response that directly follows this question: "{question}"
                LABELED TRANSCRIPT: {st.session_state.labeled_transcript}"""
                answer = get_ai_response(extract_prompt) or "No specific answer found by AI."

                eval_prompt = f"""**Task:** Evaluate the candidate's single response based on the question asked.
                **Rubric (Score 1-10):** Clarity, Correctness, Depth.
                **Rules:** Provide a score and justification for each category, an "overall_summary", and output in a valid JSON format.
                ---
                **CANDIDATE LEVEL:** {st.session_state.candidate_details['role_level']}
                **QUESTION ASKED:** {question}
                **CANDIDATE'S ANSWER TO EVALUATE:** {answer}"""
                evaluation = get_ai_response(eval_prompt, as_json=True)
                report_data.append({"question": question, "answer": answer, "evaluation": evaluation or {}})
        
        st.session_state.detailed_report = report_data
        st.session_state.status = 'report'
        st.rerun()

    # This block displays the final report.
    if st.session_state.status == 'report':
        st.subheader(f"Detailed Assessment for {st.session_state.candidate_details['name']}")

        # --- NEW, TWO-STEP DOWNLOAD LOGIC ---
        st.markdown("---")
        st.write("To download the report, first generate the file, then click download.")
        
        if st.button("Generate Word Document"):
            with st.spinner("Creating .docx file..."):
                word_data = create_word_document(st.session_state.candidate_details, st.session_state.detailed_report)
                st.session_state.word_data = word_data # Save the generated file in memory

        if st.session_state.get("word_data"):
            st.download_button(
                label="⬇️ Download as Word (.docx)",
                data=st.session_state.word_data,
                file_name=f"Interview_Report_{st.session_state.candidate_details['name']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        st.markdown("---")
        
        if st.session_state.detailed_report:
            for i, item in enumerate(st.session_state.detailed_report):
                with st.container(border=True):
                    # Display Question
                    st.markdown(f"**Q{i+1}: Interviewer:** {item['question']}")
                    
                    # Display Candidate Answer
                    st.markdown(f"**Candidate:** {item.get('answer', 'N/A')}")
                    
                    eval_data = item.get('evaluation', {})
                    summary = eval_data.get('overall_summary', 'No summary available.')
                    scores = eval_data.get('evaluation', {})
                    
                    # Display Assessment Text
                    st.markdown(f"**Assessment/Evaluation:** {summary}")
                    
                    # Calculate and Display Average Score
                    clarity = scores.get('clarity', {}).get('score', 0)
                    correctness = scores.get('correctness', {}).get('score', 0)
                    depth = scores.get('depth', {}).get('score', 0)
                    avg_score = round((clarity + correctness + depth) / 3)
                    
                    st.markdown(f"**Score:** {avg_score}/10")

                st.markdown("---") # Separator
        else:
            st.error("Could not generate detailed report.")

        st.button("Start New Interview", on_click=start_new_interview)
